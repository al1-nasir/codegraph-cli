"""Export indexed project data to DOCX format.

Two-tier export:
1. **Basic** (no LLM): structure + diagram + code listings
2. **Enhanced** (with LLM): adds AI explanations via hierarchical smart RAG

Uses ``python-docx`` for DOCX generation and embeds Mermaid diagrams
as PNG images (rendered via the Mermaid Ink online service).
"""

from __future__ import annotations

import io
import logging
import re
import time
import urllib.error
import urllib.parse
import urllib.request
from base64 import b64encode
from collections import Counter
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import typer
from rich.console import Console
from rich.progress import BarColumn, Progress, SpinnerColumn, TaskProgressColumn, TextColumn

logger = logging.getLogger(__name__)

export_app = typer.Typer(
    help="📄 Export project documentation to DOCX.",
    no_args_is_help=True,
    rich_markup_mode="rich",
    add_completion=False,
)


# ===================================================================
# DOCX generation helpers
# ===================================================================

def _ensure_docx() -> Any:
    """Import python-docx or raise a friendly error."""
    try:
        import docx
        return docx
    except ImportError:
        raise typer.Exit(
            "python-docx is required for export. Install it:\n"
            "  pip install python-docx"
        )


def _mermaid_to_png(mermaid_code: str) -> Optional[bytes]:
    """Render Mermaid code to PNG via the mermaid.ink online service.

    Falls back to None on network errors so the export can continue
    without embedded diagrams.
    """
    try:
        import base64 as _b64

        # mermaid.ink accepts base64-encoded mermaid text
        encoded = _b64.urlsafe_b64encode(mermaid_code.encode("utf-8")).decode("ascii")
        url = f"https://mermaid.ink/img/{encoded}?type=png&bgColor=!0d1117&theme=dark"

        req = urllib.request.Request(url, headers={"User-Agent": "CodeGraph-CLI/2.0"})
        with urllib.request.urlopen(req, timeout=25) as resp:
            return resp.read()
    except Exception as exc:
        logger.warning("Mermaid-to-PNG render failed: %s", exc)
        return None


def _add_mermaid_code_block(doc: Any, mermaid_code: str) -> None:
    """Add Mermaid source as a styled code block (fallback when image fails)."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(mermaid_code)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def _add_code_block(doc: Any, code: str, language: str = "python") -> None:
    """Add a monospaced code block to the DOCX document."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.style = doc.styles["No Spacing"]
    # Limit very long files
    display_code = code if len(code) < 20_000 else code[:20_000] + "\n\n# ... (truncated)"
    run = p.add_run(display_code)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xC9, 0xD1, 0xD9)


def _add_table(doc: Any, headers: List[str], rows: List[List[str]]) -> None:
    """Add a simple styled table to the DOCX."""
    from docx.shared import Pt

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Shading Accent 1"
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)


def _embed_diagram(doc: Any, mermaid_code: str, caption: str = "") -> None:
    """Attempt to embed a Mermaid diagram as a PNG image.

    Falls back to a code block with the Mermaid source if rendering fails.
    """
    from docx.shared import Inches

    png_data = _mermaid_to_png(mermaid_code)
    if png_data:
        stream = io.BytesIO(png_data)
        doc.add_picture(stream, width=Inches(6.0))
        if caption:
            p = doc.add_paragraph(caption)
            p.style = doc.styles["Caption"] if "Caption" in [s.name for s in doc.styles] else doc.styles["Normal"]
    else:
        doc.add_paragraph("(Diagram rendered as Mermaid code — paste into mermaid.live to view)")
        _add_mermaid_code_block(doc, mermaid_code)


# ===================================================================
# Data gathering from GraphStore
# ===================================================================

def _gather_project_data(store: Any) -> Dict[str, Any]:
    """Gather all project data needed for the DOCX export."""
    from .cli_explore import _build_api_tree, _generate_file_mermaid, _generate_system_mermaid

    nodes = store.get_nodes()
    edges = store.get_edges()
    nodes_by_file = store.all_by_file()
    metadata = store.get_metadata()

    # Stats
    files = sorted(nodes_by_file.keys())
    functions = [n for n in nodes if n["node_type"] == "function"]
    classes = [n for n in nodes if n["node_type"] == "class"]

    # Language breakdown by extension
    ext_counter: Counter = Counter()
    for fp in files:
        ext = Path(fp).suffix or "(no ext)"
        ext_counter[ext] += 1

    # Directory tree (reuse explore helper)
    tree = _build_api_tree(store)

    # System Mermaid diagram
    system_mermaid = _generate_system_mermaid(store)

    # Per-file Mermaid diagrams
    file_mermaids: Dict[str, str] = {}
    for fp, fnodes in nodes_by_file.items():
        diagram = _generate_file_mermaid(store, fp, fnodes)
        if diagram:
            file_mermaids[fp] = diagram

    return {
        "metadata": metadata,
        "files": files,
        "nodes_by_file": nodes_by_file,
        "functions": functions,
        "classes": classes,
        "edges": edges,
        "total_nodes": len(nodes),
        "ext_breakdown": ext_counter,
        "tree": tree,
        "system_mermaid": system_mermaid,
        "file_mermaids": file_mermaids,
    }


def _build_text_tree(node: Dict, indent: int = 0) -> str:
    """Render the JSON tree as an indented text tree."""
    lines: list[str] = []
    prefix = "  " * indent
    if node.get("type") == "dir" and node.get("name") != "root":
        lines.append(f"{prefix}📁 {node['name']}/")
    elif node.get("type") == "file":
        lines.append(f"{prefix}📄 {node['name']}")

    for child in sorted(node.get("children", []), key=lambda c: (c["type"] != "dir", c["name"])):
        lines.append(_build_text_tree(child, indent + 1))
    return "\n".join(lines)


def _get_file_deps(store: Any, file_path: str, nodes_by_file: Dict) -> Tuple[List[str], List[str]]:
    """Get dependencies and dependents for a file."""
    file_nodes = nodes_by_file.get(file_path, [])
    node_ids = {n.get("node_id") for n in file_nodes}
    deps: set[str] = set()
    dependents: set[str] = set()

    for nid in node_ids:
        for edge in store.neighbors(nid):
            dst = edge["dst"] if isinstance(edge, dict) else edge[1]
            dst_node = store.get_node(dst)
            if dst_node:
                deps.add(dst_node["qualname"])
        for edge in store.reverse_neighbors(nid):
            src = edge["src"] if isinstance(edge, dict) else edge[0]
            src_node = store.get_node(src)
            if src_node:
                dependents.add(src_node["qualname"])

    return sorted(deps), sorted(dependents)


# ===================================================================
# LLM-enhanced explanations — hierarchical strategy
# ===================================================================

def _create_llm(provider: str, model: str, api_key: str) -> Optional[Any]:
    """Create an LLM instance or return None."""
    try:
        from .llm import LocalLLM
        return LocalLLM(model=model, provider=provider, api_key=api_key)
    except Exception:
        return None


def _explain_project_overview(
    llm: Any, metadata: Dict, data: Dict, store: Any,
) -> Optional[str]:
    """Generate a project-level architecture overview via LLM."""
    # Gather compact context: directory structure + top connected files
    tree_text = _build_text_tree(data["tree"])
    # Top files by connectivity
    connectivity: list[tuple[str, int]] = []
    nodes_by_file = data["nodes_by_file"]
    for fp, fnodes in nodes_by_file.items():
        nids = {n["node_id"] for n in fnodes}
        edge_count = sum(1 for nid in nids for _ in store.neighbors(nid))
        edge_count += sum(1 for nid in nids for _ in store.reverse_neighbors(nid))
        connectivity.append((fp, edge_count))
    connectivity.sort(key=lambda x: x[1], reverse=True)
    top_files = [f for f, _ in connectivity[:10]]

    prompt = (
        "You are a senior software architect. Explain this project's "
        "architecture in 2-3 clear paragraphs. Focus on the overall design, "
        "key modules, and how they interact.\n\n"
        f"Project: {metadata.get('project_name', 'Unknown')}\n"
        f"Files: {len(data['files'])}, Functions: {len(data['functions'])}, "
        f"Classes: {len(data['classes'])}\n\n"
        f"Directory structure:\n{tree_text[:2000]}\n\n"
        f"Most connected files: {', '.join(top_files)}"
    )
    try:
        return llm.explain(prompt)
    except Exception as exc:
        logger.warning("Project overview LLM call failed: %s", exc)
        return None


def _explain_module(
    llm: Any, module_path: str, module_files: List[str],
    store: Any, nodes_by_file: Dict,
) -> Optional[str]:
    """Generate a per-module/directory summary via LLM."""
    # Gather file summaries for the module
    file_summaries: list[str] = []
    for fp in module_files[:15]:  # Limit to 15 files per module
        fnodes = nodes_by_file.get(fp, [])
        fns = [n["name"] for n in fnodes if n.get("node_type") == "function"][:10]
        clss = [n["name"] for n in fnodes if n.get("node_type") == "class"][:5]
        sig = f"  {fp}: functions={fns}, classes={clss}"
        file_summaries.append(sig)

    context = "\n".join(file_summaries)[:3000]

    prompt = (
        f"Explain the purpose of the '{module_path}' module/directory "
        f"in 2-3 sentences. What is its role and key components?\n\n"
        f"Files in module:\n{context}"
    )
    try:
        return llm.explain(prompt)
    except Exception:
        return None


def _explain_file_for_export(
    llm: Any, file_path: str, store: Any, nodes_by_file: Dict,
    source_root: str,
) -> Optional[str]:
    """Generate a per-file explanation via LLM."""
    # Read actual file if available
    content = ""
    if source_root:
        actual = Path(source_root) / file_path
        if actual.exists():
            try:
                content = actual.read_text(encoding="utf-8", errors="replace")[:4000]
            except Exception:
                pass

    if not content:
        fnodes = nodes_by_file.get(file_path, [])
        content = "\n".join(n.get("code", "")[:500] for n in fnodes[:5])

    if not content:
        return None

    deps, dependents = _get_file_deps(store, file_path, nodes_by_file)

    prompt = (
        f"Explain the purpose of '{file_path}' in 2-3 sentences. "
        f"Focus on its role and key functionality.\n\n"
        f"Dependencies: {', '.join(deps[:10]) or 'none'}\n"
        f"Dependents: {', '.join(dependents[:10]) or 'none'}\n\n"
        f"```\n{content[:3000]}\n```"
    )
    try:
        return llm.explain(prompt)
    except Exception:
        return None


# ===================================================================
# DOCX builders
# ===================================================================

def generate_basic_docx(
    store: Any,
    output_path: Path,
    include_code: bool = False,
    include_diagram: bool = True,
    console: Optional[Console] = None,
) -> Path:
    """Generate a basic DOCX without LLM — structure + diagram + code.

    Returns the path to the generated file.
    """
    docx_mod = _ensure_docx()
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    console = console or Console()

    with Progress(
        SpinnerColumn(),
        TextColumn("[progress.description]{task.description}"),
        BarColumn(),
        TaskProgressColumn(),
        console=console,
    ) as progress:
        task = progress.add_task("Gathering project data...", total=4)

        data = _gather_project_data(store)
        metadata = data["metadata"]
        progress.update(task, advance=1, description="Building document...")

        # ── Create document ──────────────────────────────────────
        doc = docx_mod.Document()

        # Title
        title = doc.add_heading(
            f"{metadata.get('project_name', 'Project')} — Documentation", level=0,
        )

        # ── Section 1: Project Overview ──────────────────────────
        doc.add_heading("1. Project Overview", level=1)

        overview_items = [
            ("Project Name", metadata.get("project_name", "Unknown")),
            ("Source Path", metadata.get("source_path", "N/A")),
            ("Total Files", str(len(data["files"]))),
            ("Total Functions", str(len(data["functions"]))),
            ("Total Classes", str(len(data["classes"]))),
            ("Total Nodes", str(data["total_nodes"])),
            ("Total Edges", str(len(data["edges"]))),
        ]
        _add_table(doc, ["Property", "Value"], overview_items)

        # Language breakdown
        if data["ext_breakdown"]:
            doc.add_heading("Language Breakdown", level=2)
            lang_rows = [(ext, str(count)) for ext, count in data["ext_breakdown"].most_common()]
            _add_table(doc, ["Extension", "File Count"], lang_rows)

        progress.update(task, advance=1, description="Adding directory tree...")

        # ── Section 2: Directory Structure ───────────────────────
        doc.add_heading("2. Directory Structure", level=1)
        tree_text = _build_text_tree(data["tree"])
        p = doc.add_paragraph()
        run = p.add_run(tree_text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

        # ── Section 3: Architecture Diagram ──────────────────────
        doc.add_heading("3. Architecture Diagram", level=1)
        doc.add_paragraph(
            "System-level dependency diagram showing file-to-file relationships."
        )
        if include_diagram and data["system_mermaid"]:
            _embed_diagram(doc, data["system_mermaid"], "System Architecture")

        progress.update(task, advance=1, description="Adding file listings...")

        # ── Section 4: File Listings ─────────────────────────────
        doc.add_heading("4. File Listings", level=1)

        source_root = metadata.get("source_path", "")

        for i, fp in enumerate(data["files"]):
            doc.add_heading(f"4.{i + 1} {fp}", level=2)

            file_nodes = data["nodes_by_file"].get(fp, [])
            fns = [n for n in file_nodes if n.get("node_type") == "function"]
            clss = [n for n in file_nodes if n.get("node_type") == "class"]

            # Symbols table
            if fns or clss:
                symbol_rows: list[list[str]] = []
                for fn in fns:
                    symbol_rows.append([
                        fn.get("name", ""),
                        "function",
                        f"L{fn.get('start_line', '?')}-{fn.get('end_line', '?')}",
                    ])
                for cls in clss:
                    symbol_rows.append([
                        cls.get("name", ""),
                        "class",
                        f"L{cls.get('start_line', '?')}-{cls.get('end_line', '?')}",
                    ])
                _add_table(doc, ["Symbol", "Type", "Lines"], symbol_rows)

            # Dependencies
            deps, dependents = _get_file_deps(store, fp, data["nodes_by_file"])
            if deps:
                doc.add_paragraph(f"Dependencies: {', '.join(deps[:20])}")
            if dependents:
                doc.add_paragraph(f"Dependents: {', '.join(dependents[:20])}")

            # Source code
            if include_code and source_root:
                actual = Path(source_root) / fp
                if actual.exists():
                    try:
                        code = actual.read_text(encoding="utf-8", errors="replace")
                        _add_code_block(doc, code)
                    except Exception:
                        doc.add_paragraph("(Could not read source file)")

            # Per-file diagram
            file_mermaid = data["file_mermaids"].get(fp)
            if include_diagram and file_mermaid:
                doc.add_heading("File Diagram", level=3)
                _embed_diagram(doc, file_mermaid, f"{Path(fp).name} internal structure")

        progress.update(task, advance=1, description="Saving...")

        doc.save(str(output_path))

    return output_path


def generate_enhanced_docx(
    store: Any,
    output_path: Path,
    include_code: bool = False,
    include_diagram: bool = True,
    explanation_depth: str = "modules",  # "overview" | "modules" | "files"
    llm_provider: str = "",
    llm_model: str = "",
    llm_api_key: str = "",
    console: Optional[Console] = None,
    batch_size: int = 5,
) -> Path:
    """Generate an LLM-enhanced DOCX with AI explanations.

    Args:
        explanation_depth: "overview" (just project), "modules" (per-dir),
                           or "files" (per-file).
        batch_size: How many LLM calls per batch (rate limiting).
    """
    docx_mod = _ensure_docx()
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    console = console or Console()

    llm = _create_llm(llm_provider, llm_model, llm_api_key)
    if not llm:
        console.print("[yellow]⚠ LLM not available — falling back to basic export.[/yellow]")
        return generate_basic_docx(store, output_path, include_code, include_diagram, console)

    data = _gather_project_data(store)
    metadata = data["metadata"]
    source_root = metadata.get("source_path", "")

    # ── Determine explanation targets ────────────────────────────
    explanation_targets: list[tuple[str, str]] = []  # (type, key)
    explanation_targets.append(("project", "project"))

    if explanation_depth in ("modules", "files"):
        # Group files by top-level directory
        modules: dict[str, list[str]] = {}
        for fp in data["files"]:
            parts = Path(fp).parts
            module_name = parts[0] if len(parts) > 1 else "(root)"
            modules.setdefault(module_name, []).append(fp)
        for mod_name in sorted(modules):
            explanation_targets.append(("module", mod_name))

    if explanation_depth == "files":
        for fp in data["files"]:
            explanation_targets.append(("file", fp))

    total_explanations = len(explanation_targets)

    # ── Generate explanations with progress ──────────────────────
    explanations: dict[str, str] = {}

    with Progress(
        SpinnerColumn(),
        TextColumn("[progress.description]{task.description}"),
        BarColumn(),
        TaskProgressColumn(),
        console=console,
    ) as progress:
        task = progress.add_task(
            f"Generating {total_explanations} AI explanations...",
            total=total_explanations,
        )

        # Group files by module for the module explanations
        modules: dict[str, list[str]] = {}
        for fp in data["files"]:
            parts = Path(fp).parts
            module_name = parts[0] if len(parts) > 1 else "(root)"
            modules.setdefault(module_name, []).append(fp)

        batch_count = 0
        for target_type, target_key in explanation_targets:
            if target_type == "project":
                progress.update(task, description="Explaining project overview...")
                result = _explain_project_overview(llm, metadata, data, store)
                if result:
                    explanations["project"] = result

            elif target_type == "module":
                progress.update(task, description=f"Explaining module: {target_key}...")
                module_files = modules.get(target_key, [])
                result = _explain_module(
                    llm, target_key, module_files, store, data["nodes_by_file"],
                )
                if result:
                    explanations[f"module:{target_key}"] = result

            elif target_type == "file":
                progress.update(task, description=f"Explaining: {target_key}...")
                result = _explain_file_for_export(
                    llm, target_key, store, data["nodes_by_file"], source_root,
                )
                if result:
                    explanations[f"file:{target_key}"] = result

            progress.update(task, advance=1)

            # Rate limiting — pause between batches
            batch_count += 1
            if batch_count % batch_size == 0 and batch_count < total_explanations:
                time.sleep(1.0)

    # ── Build the DOCX ───────────────────────────────────────────
    console.print(f"[dim]Generated {len(explanations)}/{total_explanations} explanations.[/dim]")

    with Progress(
        SpinnerColumn(),
        TextColumn("[progress.description]{task.description}"),
        BarColumn(),
        TaskProgressColumn(),
        console=console,
    ) as progress:
        task = progress.add_task("Building document...", total=5)

        doc = docx_mod.Document()

        # Title
        doc.add_heading(
            f"{metadata.get('project_name', 'Project')} — Documentation", level=0,
        )
        doc.add_paragraph("Generated by CodeGraph CLI with AI-powered explanations.")

        progress.update(task, advance=1, description="Project overview...")

        # ── Section 1: Project Overview ──────────────────────────
        doc.add_heading("1. Project Overview", level=1)

        overview_items = [
            ("Project Name", metadata.get("project_name", "Unknown")),
            ("Source Path", metadata.get("source_path", "N/A")),
            ("Total Files", str(len(data["files"]))),
            ("Total Functions", str(len(data["functions"]))),
            ("Total Classes", str(len(data["classes"]))),
            ("Total Nodes", str(data["total_nodes"])),
        ]
        _add_table(doc, ["Property", "Value"], overview_items)

        # AI overview
        if "project" in explanations:
            doc.add_heading("Architecture Overview", level=2)
            doc.add_paragraph(explanations["project"])

        # Language breakdown
        if data["ext_breakdown"]:
            doc.add_heading("Language Breakdown", level=2)
            lang_rows = [(ext, str(count)) for ext, count in data["ext_breakdown"].most_common()]
            _add_table(doc, ["Extension", "File Count"], lang_rows)

        progress.update(task, advance=1, description="Directory tree...")

        # ── Section 2: Directory Structure ───────────────────────
        doc.add_heading("2. Directory Structure", level=1)
        tree_text = _build_text_tree(data["tree"])
        p = doc.add_paragraph()
        run = p.add_run(tree_text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

        # ── Section 3: Architecture Diagram ──────────────────────
        doc.add_heading("3. Architecture Diagram", level=1)
        doc.add_paragraph(
            "System-level dependency diagram showing file-to-file relationships."
        )
        if include_diagram and data["system_mermaid"]:
            _embed_diagram(doc, data["system_mermaid"], "System Architecture")

        progress.update(task, advance=1, description="Module summaries...")

        # ── Section 4: Module Summaries ──────────────────────────
        if explanation_depth in ("modules", "files"):
            doc.add_heading("4. Module Summaries", level=1)

            modules_sorted = sorted(modules.items())
            for mod_name, mod_files in modules_sorted:
                doc.add_heading(f"📁 {mod_name}", level=2)

                # Module explanation
                mod_key = f"module:{mod_name}"
                if mod_key in explanations:
                    doc.add_paragraph(explanations[mod_key])

                # File list for this module
                file_rows = []
                for fp in mod_files:
                    fnodes = data["nodes_by_file"].get(fp, [])
                    fn_count = sum(1 for n in fnodes if n.get("node_type") == "function")
                    cls_count = sum(1 for n in fnodes if n.get("node_type") == "class")
                    file_rows.append([Path(fp).name, str(fn_count), str(cls_count)])
                if file_rows:
                    _add_table(doc, ["File", "Functions", "Classes"], file_rows)

            section_offset = 5
        else:
            section_offset = 4

        progress.update(task, advance=1, description="File listings...")

        # ── Section N: File Listings ─────────────────────────────
        doc.add_heading(f"{section_offset}. File Listings", level=1)

        for i, fp in enumerate(data["files"]):
            doc.add_heading(f"{section_offset}.{i + 1} {fp}", level=2)

            # AI file explanation
            file_key = f"file:{fp}"
            if file_key in explanations:
                p = doc.add_paragraph()
                run = p.add_run("🤖 AI Explanation: ")
                run.bold = True
                run.font.color.rgb = RGBColor(0xBC, 0x8C, 0xFF)
                p.add_run(explanations[file_key])

            file_nodes = data["nodes_by_file"].get(fp, [])
            fns = [n for n in file_nodes if n.get("node_type") == "function"]
            clss = [n for n in file_nodes if n.get("node_type") == "class"]

            if fns or clss:
                symbol_rows: list[list[str]] = []
                for fn in fns:
                    symbol_rows.append([
                        fn.get("name", ""),
                        "function",
                        f"L{fn.get('start_line', '?')}-{fn.get('end_line', '?')}",
                    ])
                for cls in clss:
                    symbol_rows.append([
                        cls.get("name", ""),
                        "class",
                        f"L{cls.get('start_line', '?')}-{cls.get('end_line', '?')}",
                    ])
                _add_table(doc, ["Symbol", "Type", "Lines"], symbol_rows)

            # Dependencies
            deps, dependents = _get_file_deps(store, fp, data["nodes_by_file"])
            if deps:
                doc.add_paragraph(f"Dependencies: {', '.join(deps[:20])}")
            if dependents:
                doc.add_paragraph(f"Dependents: {', '.join(dependents[:20])}")

            # Source code
            if include_code and source_root:
                actual = Path(source_root) / fp
                if actual.exists():
                    try:
                        code = actual.read_text(encoding="utf-8", errors="replace")
                        _add_code_block(doc, code)
                    except Exception:
                        doc.add_paragraph("(Could not read source file)")

            # Per-file diagram
            file_mermaid = data["file_mermaids"].get(fp)
            if include_diagram and file_mermaid:
                doc.add_heading("File Diagram", level=3)
                _embed_diagram(doc, file_mermaid, f"{Path(fp).name} internal structure")

        progress.update(task, advance=1, description="Saving...")

        doc.save(str(output_path))

    return output_path


# ===================================================================
# JSON export for browser integration
# ===================================================================

def generate_export_data(
    store: Any,
    include_code: bool = False,
    llm_provider: str = "",
    llm_model: str = "",
    llm_api_key: str = "",
    explanation_depth: str = "overview",
) -> Dict[str, Any]:
    """Generate export data as JSON (used by the browser UI).

    Returns a dict suitable for JSON serialization with all sections.
    """
    data = _gather_project_data(store)
    metadata = data["metadata"]
    source_root = metadata.get("source_path", "")

    result: Dict[str, Any] = {
        "project_name": metadata.get("project_name", "Unknown"),
        "source_path": source_root,
        "stats": {
            "files": len(data["files"]),
            "functions": len(data["functions"]),
            "classes": len(data["classes"]),
            "nodes": data["total_nodes"],
            "edges": len(data["edges"]),
        },
        "ext_breakdown": dict(data["ext_breakdown"]),
        "tree_text": _build_text_tree(data["tree"]),
        "system_mermaid": data["system_mermaid"],
        "files": [],
    }

    for fp in data["files"]:
        file_nodes = data["nodes_by_file"].get(fp, [])
        fns = [{"name": n["name"], "line": n["start_line"]}
               for n in file_nodes if n.get("node_type") == "function"]
        clss = [{"name": n["name"], "line": n["start_line"]}
                for n in file_nodes if n.get("node_type") == "class"]
        deps, dependents = _get_file_deps(store, fp, data["nodes_by_file"])

        file_entry: Dict[str, Any] = {
            "path": fp,
            "name": Path(fp).name,
            "functions": fns,
            "classes": clss,
            "deps": deps[:20],
            "dependents": dependents[:20],
        }

        if include_code and source_root:
            actual = Path(source_root) / fp
            if actual.exists():
                try:
                    file_entry["content"] = actual.read_text(
                        encoding="utf-8", errors="replace",
                    )
                except Exception:
                    pass

        result["files"].append(file_entry)

    return result


# ===================================================================
# Typer commands
# ===================================================================

@export_app.command("docx")
def export_docx(
    output: Path = typer.Option(
        None, "--output", "-o",
        help="Output file path (default: <project>.docx)",
    ),
    include_code: bool = typer.Option(
        False, "--include-code", "-c",
        help="Include source code in the export.",
    ),
    enhanced: bool = typer.Option(
        False, "--enhanced", "-e",
        help="Use LLM for AI-powered explanations (requires configured LLM).",
    ),
    depth: str = typer.Option(
        "modules", "--depth", "-d",
        help="Explanation depth: overview, modules, or files.",
    ),
    no_diagram: bool = typer.Option(
        False, "--no-diagram",
        help="Skip embedding architecture diagram (faster, smaller file).",
    ),
):
    """📄 Export project documentation to DOCX format.

    [bold]Basic export[/bold] (default):
      Structure, directory tree, architecture diagram, file symbols

    [bold]Enhanced export[/bold] (--enhanced):
      Adds AI explanations at chosen depth level

    [bold]Examples:[/bold]
      cg export docx
      cg export docx --output my-docs.docx --include-code
      cg export docx --enhanced --depth files --include-code
    """
    from .storage import GraphStore, ProjectManager
    from . import config as cfg

    console = Console()

    pm = ProjectManager()
    project = pm.get_current_project()
    if not project:
        console.print("[red]No project loaded.[/red] Run [cyan]cg project index <path>[/cyan] first.")
        raise typer.Exit(code=1)

    project_dir = pm.project_dir(project)
    if not project_dir.exists():
        console.print(f"[red]Project '{project}' not found.[/red]")
        raise typer.Exit(code=1)

    store = GraphStore(project_dir)

    nodes = store.get_nodes()
    if not nodes:
        console.print("[red]Project is empty.[/red] Re-run [cyan]cg project index[/cyan].")
        store.close()
        raise typer.Exit(code=1)

    # Default output path
    if output is None:
        output = Path(f"{project}-docs.docx")

    # Validate depth
    if depth not in ("overview", "modules", "files"):
        console.print(f"[red]Invalid depth '{depth}'. Use: overview, modules, or files.[/red]")
        store.close()
        raise typer.Exit(code=1)

    try:
        console.print(f"\n[bold green]📄 Exporting to DOCX[/bold green]")
        console.print(f"   Project:  [cyan]{project}[/cyan]")
        console.print(f"   Output:   [cyan]{output}[/cyan]")
        console.print(f"   Code:     {'✅' if include_code else '❌'}")
        console.print(f"   Diagram:  {'✅' if not no_diagram else '❌'}")
        console.print(f"   Enhanced: {'✅ ' + depth if enhanced else '❌'}")
        console.print()

        if enhanced:
            result = generate_enhanced_docx(
                store=store,
                output_path=output,
                include_code=include_code,
                include_diagram=not no_diagram,
                explanation_depth=depth,
                llm_provider=cfg.LLM_PROVIDER,
                llm_model=cfg.LLM_MODEL,
                llm_api_key=cfg.LLM_API_KEY,
                console=console,
            )
        else:
            result = generate_basic_docx(
                store=store,
                output_path=output,
                include_code=include_code,
                include_diagram=not no_diagram,
                console=console,
            )

        file_size = output.stat().st_size
        size_str = (
            f"{file_size / 1024 / 1024:.1f} MB" if file_size > 1024 * 1024
            else f"{file_size / 1024:.1f} KB"
        )
        console.print(f"\n[bold green]✅ Exported successfully![/bold green]")
        console.print(f"   File: [cyan]{result}[/cyan] ({size_str})")

    except Exception as e:
        console.print(f"\n[red]Export failed: {e}[/red]")
        logger.exception("Export failed")
        raise typer.Exit(code=1)
    finally:
        store.close()
